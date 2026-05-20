"""
Build .docx per language for the TOF 17.05 campaign.
Each ad: filename heading -> embedded image -> 3 headlines -> primary text -> CTA suggestion.
Saved into the same Vein new ads 17.05/ folder so they live next to the creatives.
"""
import io
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

FOLDER = Path("/Users/magi/Downloads/Vein new ads 17.05")

LANGUAGES = [
    ("EN", "English",                "UK, US, AU, IE",                 "inaessentials.co.uk, inaessentials.us, inaessentials.au, inaessentials.ie"),
    ("BG", "Bulgarian",              "BG",                              "inaessentials.bg"),
    ("FR", "French",                 "FR",                              "inaessentials.fr"),
    ("RO", "Romanian",               "RO",                              "inaessentials.ro"),
    ("SK", "Slovak",                 "SK",                              "www.inaessentials.sk"),
    ("CZ", "Czech",                  "CZ",                              "www.inaessentials.cz"),
    ("DE", "German",                 "DE",                              "inaessentials.de"),
    ("IT", "Italian",                "IT",                              "inaessentials.it"),
    ("ES", "Spanish",                "ES",                              "inaessentials.es"),
    ("NL", "Dutch",                  "NL",                              "inaessentials.nl"),
    ("PT", "Portuguese (European)",  "PT",                              "inaessentials.pt"),
    ("PL", "Polish",                 "PL",                              "inaessentials.pl"),
    ("HU", "Hungarian",              "HU",                              "inaessentials.hu"),
    ("HR", "Croatian",               "HR",                              "inaessentials.hr"),
    ("SI", "Slovenian",              "SI",                              "inaessentials.si"),
    ("RS", "Serbian (Latin)",        "RS",                              "inaessentials.rs"),
    ("EL", "Greek",                  "GR",                              "inaessentials.gr"),
    ("DK", "Danish",                 "DK",                              "inaessentials.dk"),
    ("LT", "Lithuanian",             "LT",                              "inaessentials.lt"),
    ("LV", "Latvian",                "LV",                              "inaessentials.lv"),
    ("SE", "Swedish",                "SE",                              "inaessentials.se"),
]

IMG_MAX_WIDTH_INCHES = 4.0
IMG_MAX_DIM_PX = 1200


def prepare_image_bytes(src_path: Path):
    with Image.open(src_path) as im:
        im.load()
        if im.mode in ("RGBA", "P"):
            bg = Image.new("RGB", im.size, (255, 255, 255))
            if im.mode == "P":
                im = im.convert("RGBA")
            bg.paste(im, mask=im.split()[-1] if im.mode == "RGBA" else None)
            im = bg
        elif im.mode != "RGB":
            im = im.convert("RGB")
        w, h = im.size
        scale = min(1.0, IMG_MAX_DIM_PX / max(w, h))
        if scale < 1.0:
            im = im.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
        buf = io.BytesIO()
        im.save(buf, format="JPEG", quality=82, optimize=True)
        buf.seek(0)
        return buf


def add_kv(doc, label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    p.add_run(value)


def build_doc(code, name, markets, domains, data, out_path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Veins TOF, 36 ads, 17.05", level=0)
    doc.add_heading(f"Ad copy, {name}", level=1)
    add_kv(doc, "Markets", markets)
    add_kv(doc, "Shopify domains", domains)
    add_kv(doc, "Total ads", str(len(data)))
    add_kv(doc, "Format", "Image, then 3 headlines, then primary text, then CTA suggestion")
    add_kv(doc, "Tone", "Top of funnel, problem-aware, solution-unaware. No prices, no product name.")
    doc.add_paragraph()

    for i, r in enumerate(data, 1):
        fn = r["filename"]
        doc.add_heading(f"{i}. {fn}", level=2)
        img_path = FOLDER / fn
        if img_path.exists():
            try:
                buf = prepare_image_bytes(img_path)
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(buf, width=Inches(IMG_MAX_WIDTH_INCHES))
            except Exception as e:
                doc.add_paragraph(f"[image error: {e}]")
        else:
            doc.add_paragraph(f"[image not found: {fn}]")

        p = doc.add_paragraph()
        p.add_run("Headlines").bold = True
        for j, h in enumerate(r.get("headlines", []), 1):
            doc.add_paragraph(f"{j}. {h}")
        p2 = doc.add_paragraph()
        p2.add_run("Primary Text").bold = True
        doc.add_paragraph(r.get("primary_text", ""))
        add_kv(doc, "CTA suggestion", r.get("cta_suggestion", "LEARN_MORE"))
        sep = doc.add_paragraph("_" * 50)
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(out_path))


def main():
    for code, name, markets, domains in LANGUAGES:
        src = FOLDER / f"ad_copy_TOF_{code}.json"
        if not src.exists():
            print(f"[{code}] no translation file, skipping")
            continue
        data = json.loads(src.read_text())
        slug = name.replace(" ", "_").replace("(", "").replace(")", "")
        out = FOLDER / f"Veins_TOF_{code}_{slug}.docx"
        build_doc(code, name, markets, domains, data, out)
        size_kb = out.stat().st_size // 1024
        print(f"[{code}] wrote {out.name} ({size_kb} KB)")


if __name__ == "__main__":
    main()
