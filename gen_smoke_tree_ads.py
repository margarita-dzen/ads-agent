"""
Generate Bulgarian ad copy for the AI ADS MAY/ folders using the same
SYSTEM_PROMPT + generate_copy logic as the local Ad Copy Agent
(pages/1_✦_Ad_Copy.py). Builds one .docx per folder with each image
embedded above its Headlines / Primary Text / CTA block.

Brand: Ina Essentials Smoke Tree Hydrolina (oral spray, Cotinus coggygria).
Voice: Bulgarian, evergreen, problem-aware, informal "ти", NO em-dashes.
"""
import base64
import io
import os
import random
import sys
import time
from pathlib import Path

import anthropic
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from PIL import Image

load_dotenv(Path(__file__).parent / ".env")

# ── Folders ─────────────────────────────────────────────────────────────────
ROOT = Path("/Users/magi/Downloads/AI ADS MAY")
FOLDERS = [
    ("General", "General_AdCopy_BG.docx", "General creatives, broader top-of-funnel hooks"),
    ("New Concept (GUM RECCESION)", "GumRecession_AdCopy_BG.docx",
     "Diagnostic close-ups for gum recession / bleeding gums / gingivitis / tartar / plaque"),
]

# ── Same SYSTEM_PROMPT as pages/1_✦_Ad_Copy.py ──────────────────────────────
SYSTEM_PROMPT = """You are an expert direct-response advertising copywriter with 15+ years of experience across Meta, Google, TikTok, and display advertising.

Respond in this EXACT format:

AD ANGLE
[one sentence naming the angle and why it fits this visual]

HEADLINES
1. [headline]
2. [headline]
3. [headline]

PRIMARY TEXT
[the copy — follow the length instruction exactly]

CTA
[call-to-action button text only]

Rules:
- No generic copy. Every line must feel specific to THIS image.
- Headlines under 10 words each.
- If a product/brand/offer is visible, use it.
- Match tone to the visual energy."""

COPY_LENGTH_INSTRUCTIONS = {
    "Short":  "Primary Text: 1–2 sentences. Ultra-concise — hook immediately, every word earns its place.",
    "Medium": "Primary Text: 3–4 sentences. Hook + context + payoff. Balanced for standard feed placements.",
    "Long":   "Primary Text: 6–10 sentences. Story arc: strong hook, build desire, present product as solution, close with urgency or social proof.",
}

PRODUCT_INFO = """Product: Хидролина Скумпия (Ina Essentials Smoke Tree Hydrolina)
Type: Natural oral care spray, Bulgarian herbal heritage product.
Active ingredient: Cotinus coggygria (Скумпия / Смрадлика), a Balkan herb traditionally used for centuries against bleeding gums, gingivitis, and bad breath.
Use cases: bleeding gums, sensitive gums, gum recession risk, persistent bad breath, oral microbiome support, mouth ulcers, after dental cleaning.
Brand: Ina Essentials — Bulgarian natural cosmetics, Hydrolina line.
Packaging: small spray bottle, red cap, red/orange label.
Positioning: forgotten Bulgarian remedy in modern spray format. Not a medicinal product, no clinical claims. Pure herbal hydrolat.
Tone of brand: warm, heritage-driven, calm, knowledgeable, "ти" informal Bulgarian."""

BG_CONTEXT_BASE = """LANGUAGE & STYLE REQUIREMENTS (CRITICAL):
- Write ALL output in Bulgarian (Cyrillic). Do not output a single line in English.
- AD ANGLE itself stays in English (it's a strategy label for the marketer), but HEADLINES, PRIMARY TEXT and CTA MUST be Bulgarian Cyrillic.
- Use informal "ти" form, never "вие/Вие".
- DO NOT use em-dashes (—) or en-dashes (–) anywhere. Use commas, periods, colons or new sentences instead. This is a hard brand rule.
- Evergreen tone, no fake urgency, no "limited time", no discount/price mentions.
- CTA must be a real Meta button label in Bulgarian: "Научи повече", "Виж повече", "Поръчай сега", "Купи сега", or "Открий продукта". Pick whichever fits best.
- Primary text should: hook on the problem the IMAGE shows, name the Bulgarian herb (Скумпия / смрадлика / Cotinus coggygria), mention Хидролина as the modern spray format, end with a soft CTA sentence.
- If the image already has a Bulgarian headline baked in, your HEADLINES must COMPLEMENT it (different angle), never duplicate it word-for-word.
- Headlines under 40 characters each ideally, never above 10 words.
- No emojis."""

# ── Image embedding helper (matches build_docx.py) ──────────────────────────
IMG_MAX_DIM_PX = 1200
IMG_MAX_WIDTH_INCHES = 5.0


def prepare_image_bytes(src_path: Path) -> io.BytesIO:
    with Image.open(src_path) as im:
        im.load()
        if im.mode in ("RGBA", "P"):
            bg = Image.new("RGB", im.size, (255, 255, 255))
            if im.mode == "P":
                im = im.convert("RGBA")
            bg.paste(im, mask=im.split()[-1] if im.mode == "RGBA" else None)
            im = bg
        elif im.mode != "RGB":
            im = im.convert("RGB")
        w, h = im.size
        scale = min(1.0, IMG_MAX_DIM_PX / max(w, h))
        if scale < 1.0:
            im = im.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
        buf = io.BytesIO()
        im.save(buf, format="JPEG", quality=82, optimize=True)
        buf.seek(0)
        return buf


def compress_image_for_api(image_bytes: bytes, media_type: str,
                           max_bytes: int = 4 * 1024 * 1024) -> tuple[bytes, str]:
    if len(image_bytes) <= max_bytes:
        return image_bytes, media_type
    img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    quality = 85
    scale = 1.0
    for _ in range(8):
        w, h = int(img.width * scale), int(img.height * scale)
        resized = img.resize((w, h), Image.LANCZOS) if scale < 1.0 else img
        buf = io.BytesIO()
        resized.save(buf, format="JPEG", quality=quality)
        if buf.tell() <= max_bytes:
            return buf.getvalue(), "image/jpeg"
        scale *= 0.85
        quality = max(quality - 5, 50)
    buf = io.BytesIO()
    img.resize((800, int(800 * img.height / img.width)), Image.LANCZOS).save(buf, format="JPEG", quality=60)
    return buf.getvalue(), "image/jpeg"


# ── generate_copy + parse_result (same as agent) ────────────────────────────
def generate_copy(client, image_bytes: bytes, media_type: str, context: str) -> str:
    image_bytes, media_type = compress_image_for_api(image_bytes, media_type)
    b64 = base64.standard_b64encode(image_bytes).decode("utf-8")
    prompt = (
        "Write ad copy for this image.\n"
        "Target platform: Instagram / Facebook.\n"
        f"Length instruction: {COPY_LENGTH_INSTRUCTIONS['Medium']}\n\n"
        f"Product page information (use this to ground the copy in real product details):\n{PRODUCT_INFO}\n\n"
        f"Extra context: {context}"
    )
    user_content = [
        {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": b64}},
        {"type": "text", "text": prompt},
    ]
    for attempt in range(5):
        try:
            r = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=2048,
                system=SYSTEM_PROMPT,
                messages=[{"role": "user", "content": user_content}],
            )
            return r.content[0].text
        except anthropic.APIStatusError as e:
            if e.status_code in (529, 500) and attempt < 4:
                time.sleep(min(2 ** attempt + random.uniform(0, 1), 30))
            else:
                raise


def parse_result(text: str) -> dict:
    out = {"angle": "", "headlines": [], "primary": "", "cta": ""}
    current = None
    for line in text.strip().splitlines():
        s = line.strip()
        if s == "AD ANGLE":
            current = "angle"
        elif s == "HEADLINES":
            current = "headlines"
        elif s == "PRIMARY TEXT":
            current = "primary"
        elif s == "CTA":
            current = "cta"
        elif s and current:
            if current == "angle" and not out["angle"]:
                out["angle"] = s
            elif current == "headlines":
                clean = s.lstrip("123456789.-) ").strip()
                if clean:
                    out["headlines"].append(clean)
            elif current == "primary":
                out["primary"] += (" " if out["primary"] else "") + s
            elif current == "cta" and not out["cta"]:
                out["cta"] = s
    return out


def strip_dashes(text: str) -> str:
    if not text:
        return text
    return (text.replace("—", ", ").replace("–", ", ")
            .replace(" ,", ",").replace(",,", ","))


# ── Per-folder context hints (image-baked headlines from prompt files) ──────
IMAGE_HINTS = {
    # Gum Recession folder (ANGLE A diagnostic close-ups, OralCare_Nordics_Adapted_v4.md)
    "A01_gum_recession_diagnostic.png": "Image-baked headline: 'Знаеш ли какво виждаш? Това е РЕЦЕСИЯ.' Topic: gum recession.",
    "A02_bleeding_gums_diagnostic.png": "Image-baked headline: 'Кървят ли венците, когато миеш зъби?' Topic: early bleeding gums.",
    "A03_gingivitis_inflammation.png": "Image-baked headline: 'Червени и подути венци? Първи сигнал.' Topic: gingivitis inflammation.",
    "A04_tartar_at_gumline.png": "Image-baked headline: 'Жълтото в основата на зъбите не е от кафето.' Topic: tartar at gumline.",
    "A05_tongue_coating.png": "Image-baked headline: '80% от лошия дъх започва от ТУК.' Topic: bacterial coating on tongue, bad breath.",
    "A06_exposed_tooth_root.png": "Image-baked headline: 'Защо зъбите ти изглеждат по-дълги?' Topic: exposed tooth root from recession.",
    "A07_periodontal_pocket.png": "Image-baked headline: 'Между зъба и венеца има тих враг.' Topic: periodontal pocket.",
    "A08_plaque_buildup.png": "Image-baked headline: 'Това, което не виждаш сутрин в огледалото.' Topic: plaque buildup near gumline.",
    # Smoke tree environmental shots (Higgsfield prompts file)
    "hf_20260514_142022_79e54613-3c54-4a3e-b21a-7362b8ec9547.png":
        "Image: cinematic smoke tree (Скумпия / Cotinus coggygria) nature scene, no product visible. Use this to introduce the Bulgarian herb tradition.",
    "hf_20260514_142124_10609a2d-6875-4003-bc2b-749cd629f0d4.png":
        "Image: cinematic smoke tree nature scene, no product visible. Same herb-heritage angle.",
    "hf_20260515_060236_d5319956-8e84-4edf-95c2-91c8162495a9.png":
        "Image: cinematic smoke tree nature scene, no product visible. Same herb-heritage angle.",
    "hf_20260515_060324_c4173df7-fec4-45b6-a768-7192c042feb1.png":
        "Image: cinematic smoke tree nature scene, no product visible. Same herb-heritage angle.",
    "hf_20260515_060350_10be5a72-7baa-46d4-b014-c82b652cb80d.png":
        "Image: cinematic smoke tree nature scene, no product visible. Same herb-heritage angle.",

    # General folder (OralCare_TOF_Prompts_v3_GPT.md + v2 mix)
    "01_breaking_news_gum_headline.png":
        "Image: faux magazine cover with vintage photo of Bulgarian grandmother + headline 'Какво ползваше баба ти за устата, преди да съществуваше пастата за зъби'. Heritage angle.",
    "03_mystery_hook_article.png":
        "Image: editorial doctor's handwritten notepad with line 'Кървене. Възпаление. Лош дъх. Всичко започва от едно място.' Authority / mystery angle.",
    "04_whiteboard_why_bad_breath_returns.png":
        "Image: whiteboard-style educational graphic explaining why bad breath keeps returning. Educational angle.",
    "05_anatomy_diagram_tongue_bacteria.png":
        "Image: anatomical diagram of tongue showing bacterial coating. Educational / problem-aware.",
    "07_stat_surround_one_in_three.png":
        "Image: 'one in three' statistic surround poster about gum health. Stat-hook angle.",
    "08_ugc_mirror_teeth_check.png":
        "Image: UGC mirror selfie checking teeth/gums. Self-diagnosis curiosity angle.",
    "09_ugc_hand_covering_mouth.png":
        "Image: UGC person covering mouth, embarrassed by breath. Social-shame relatable angle.",
    "10_ugc_toothbrush_moment.png":
        "Image: UGC toothbrush moment, possibly showing pink in sink. Relatable problem moment.",
    "11_pink_smoke_tree_with_headline.png":
        "Image: pink smoke tree (Скумпия) bloom with headline overlay. Beauty + heritage angle.",
    "13_leaf_macro_with_stat.png":
        "Image: smoke tree leaf macro with a statistic overlay. Educational beauty angle.",
    "14_color_block_bold_statement.png":
        "Image: bold typographic color-block statement poster ('Лош дъх не се измива.' or similar). Provocative angle.",
}


# ── Doc builder ─────────────────────────────────────────────────────────────
def add_label(doc, text, bold=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p


def build_folder_docx(folder: Path, docx_name: str, folder_desc: str, results: list[dict]):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Ina Essentials, Хидролина Скумпия", level=0)
    doc.add_heading(f"Ad copy, Bulgarian, {folder.name}", level=1)

    meta = doc.add_paragraph()
    meta.add_run("Folder: ").bold = True
    meta.add_run(folder_desc)

    p = doc.add_paragraph()
    p.add_run("Total ads: ").bold = True
    p.add_run(str(len(results)))

    p = doc.add_paragraph()
    p.add_run("Format: ").bold = True
    p.add_run("3 заглавия (опции), основен текст, CTA. Без тирета. Език: български.")

    doc.add_paragraph()

    for r in results:
        doc.add_page_break()
        doc.add_heading(r["filename"], level=2)

        img_path = folder / r["filename"]
        if img_path.exists():
            try:
                buf = prepare_image_bytes(img_path)
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(buf, width=Inches(IMG_MAX_WIDTH_INCHES))
            except Exception as e:
                doc.add_paragraph(f"[image error: {e}]")

        if r.get("angle"):
            p = doc.add_paragraph()
            p.add_run("Ad Angle: ").bold = True
            p.add_run(strip_dashes(r["angle"]))

        add_label(doc, "Заглавия (3 опции):")
        for i, h in enumerate(r.get("headlines", []), 1):
            doc.add_paragraph(f"{i}. {strip_dashes(h)}")

        add_label(doc, "Основен текст:")
        doc.add_paragraph(strip_dashes(r.get("primary", "")))

        p = doc.add_paragraph()
        p.add_run("CTA: ").bold = True
        p.add_run(strip_dashes(r.get("cta", "")))

        sep = doc.add_paragraph("_" * 50)
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

    out = folder / docx_name
    doc.save(str(out))
    return out


# ── Main ────────────────────────────────────────────────────────────────────
def main():
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        print("ERROR: ANTHROPIC_API_KEY missing in /Users/magi/ads-agent/.env")
        sys.exit(1)
    client = anthropic.Anthropic(api_key=api_key)

    for sub, docx_name, folder_desc in FOLDERS:
        folder = ROOT / sub
        if not folder.is_dir():
            print(f"[SKIP] not a directory: {folder}")
            continue
        pngs = sorted(p for p in folder.iterdir()
                      if p.is_file() and p.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp"})
        print(f"\n=== {folder.name}: {len(pngs)} images ===")

        results = []
        for i, img_path in enumerate(pngs, 1):
            print(f"  [{i}/{len(pngs)}] {img_path.name}", flush=True)
            hint = IMAGE_HINTS.get(img_path.name, "")
            extra = BG_CONTEXT_BASE + ("\n\nIMAGE-SPECIFIC HINT:\n" + hint if hint else "")
            try:
                ext = img_path.suffix.lower().lstrip(".")
                media_type = {"jpg": "image/jpeg", "jpeg": "image/jpeg",
                              "png": "image/png", "webp": "image/webp"}.get(ext, "image/jpeg")
                raw = generate_copy(client, img_path.read_bytes(), media_type, extra)
                parsed = parse_result(raw)
                parsed["filename"] = img_path.name
                parsed["raw"] = raw
                results.append(parsed)
            except Exception as e:
                print(f"      ERROR: {e}")
                results.append({"filename": img_path.name, "angle": "", "headlines": [],
                                "primary": f"[generation error: {e}]", "cta": "", "raw": ""})

        out = build_folder_docx(folder, docx_name, folder_desc, results)
        print(f"  -> wrote {out} ({out.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
