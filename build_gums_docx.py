"""
Build per-language .docx for GUMS campaign with theme-ordered ads.
For each language, use images from that language's subfolder (matching by basename).
"""
import io
import json
import re
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

FOLDER = Path("/Users/magi/Downloads/GUMS 19.06.2026")
COPY_DIR = FOLDER / "Ad copy Sumac 19.05. "
if not COPY_DIR.exists():
    COPY_DIR = FOLDER  # fallback if folder name differs
THEME_MAP = json.loads((FOLDER / "theme_map.json").read_text())

# (code, name, folder_name, markets, domains)
LANGUAGES = [
    ("EN", "English",                "english",  "UK, US, AU, IE", "inaessentials.co.uk"),
    ("BG", "Bulgarian",              "Bulgarian","BG", "inaessentials.bg"),
    ("FR", "French",                 "FR",       "FR", "inaessentials.fr"),
    ("RO", "Romanian",               "RO",       "RO", "inaessentials.ro"),
    ("CZ", "Czech",                  "CZ",       "CZ", "www.inaessentials.cz"),
    ("DE", "German",                 "DE",       "DE", "inaessentials.de"),
    ("IT", "Italian",                "IT",       "IT", "inaessentials.it"),
    ("ES", "Spanish",                "ES",       "ES", "inaessentials.es"),
    ("NL", "Dutch",                  "NL",       "NL", "inaessentials.nl"),
    ("PT", "Portuguese (European)",  "PT",       "PT", "inaessentials.pt"),
    ("PL", "Polish",                 "PL",       "PL", "inaessentials.pl"),
    ("HU", "Hungarian",              "HU",       "HU", "inaessentials.hu"),
    ("HR", "Croatian",               "HR",       "HR", "inaessentials.hr"),
    ("RS", "Serbian (Latin)",        "RS",       "RS", "inaessentials.rs"),
    ("DK", "Danish",                 "DK",       "DK", "inaessentials.dk"),
    ("SE", "Swedish",                "SE",       "SE", "inaessentials.se"),
    ("SK", "Slovak",                 "SK",       "SK", "www.inaessentials.sk"),
    ("SI", "Slovenian",              "SI",       "SI", "inaessentials.si"),
    ("LT", "Lithuanian",             "LT",       "LT", "inaessentials.lt"),
    ("LV", "Latvian",                "LV",       "LV", "inaessentials.lv"),
]

IMG_MAX_WIDTH_INCHES = 4.0
IMG_MAX_DIM_PX = 1200


def prepare_image_bytes(src: Path):
    with Image.open(src) as im:
        im.load()
        if im.mode in ("RGBA", "P"):
            bg = Image.new("RGB", im.size, (255, 255, 255))
            if im.mode == "P":
                im = im.convert("RGBA")
            bg.paste(im, mask=im.split()[-1] if im.mode == "RGBA" else None)
            im = bg
        elif im.mode != "RGB":
            im = im.convert("RGB")
        w, h = im.size
        scale = min(1.0, IMG_MAX_DIM_PX / max(w, h))
        if scale < 1.0:
            im = im.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
        buf = io.BytesIO()
        im.save(buf, format="JPEG", quality=82, optimize=True)
        buf.seek(0)
        return buf


def add_kv(doc, label, value):
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(value)


LANG_SUFFIXES = ('_BG','_CZ','_DE','_DK','_ES','_FR','_HR','_HU','_IT','_LT','_LV','_NL','_PL','_PT','_RO','_RS','_SE','_SI','_SK','_EN','_EL','_GR')


def base_id(filename: str) -> str:
    """Strip language suffix + extension. Returns canonical creative ID."""
    stem, _, _ = filename.rpartition(".")
    if not stem:
        stem = filename
    for suf in LANG_SUFFIXES:
        if stem.endswith(suf):
            return stem[:-len(suf)]
    return stem


def find_image_path(en_filename: str, lang_folder: Path, lang_code: str) -> Path:
    """Find the localized image whose base ID matches the EN entry."""
    target_base = base_id(en_filename)
    for p in lang_folder.iterdir():
        if p.suffix.lower() not in (".png", ".jpg", ".jpeg", ".webp", ".gif"):
            continue
        if base_id(p.name) == target_base:
            return p
    return None


def build_doc(code, name, lang_folder, markets, domains, data, themes, out_path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("GUMS / Oral TOF — Smoke Tree Hydrolina campaign", level=0)
    doc.add_heading(f"Ad copy, {name}", level=1)
    add_kv(doc, "Markets", markets)
    add_kv(doc, "Source folder", str(lang_folder))
    add_kv(doc, "Tone", "Top of funnel, problem-aware (bad breath, bleeding gums, sensitivity, recession). No prices, no brand.")

    # Group entries by theme, in theme order
    by_file = {r["filename"]: r for r in data}
    by_theme = OrderedDict()
    for theme in themes["themes"]:
        by_theme[theme] = []
    for fn, theme in themes["assignments"].items():
        if fn in by_file:
            by_theme.setdefault(theme, []).append(by_file[fn])

    add_kv(doc, "Total ads", str(sum(len(v) for v in by_theme.values())))
    doc.add_paragraph()

    ad_n = 0
    skipped = 0
    for theme in themes["themes"]:
        ads = by_theme.get(theme, [])
        if not ads:
            continue
        doc.add_heading(f"Theme: {theme}  ({len(ads)} ads)", level=1)
        for r in ads:
            ad_n += 1
            fn = r["filename"]
            img_path = find_image_path(fn, lang_folder, code)
            doc.add_heading(f"{ad_n}. {fn}", level=2)
            if img_path:
                try:
                    buf = prepare_image_bytes(img_path)
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.add_run().add_picture(buf, width=Inches(IMG_MAX_WIDTH_INCHES))
                except Exception as e:
                    doc.add_paragraph(f"[image error: {e}]")
            else:
                doc.add_paragraph(f"[no localized image found in {lang_folder.name}/, skipping image]")
                skipped += 1
            p = doc.add_paragraph()
            p.add_run("Headlines").bold = True
            for j, h in enumerate(r.get("headlines", []), 1):
                doc.add_paragraph(f"{j}. {h}")
            p2 = doc.add_paragraph()
            p2.add_run("Primary Text").bold = True
            doc.add_paragraph(r.get("primary_text", ""))
            add_kv(doc, "CTA suggestion", r.get("cta_suggestion", "LEARN_MORE"))
            sep = doc.add_paragraph("_" * 50)
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(out_path))
    return ad_n, skipped


def main():
    for code, name, folder_name, markets, domains in LANGUAGES:
        src = COPY_DIR / f"ad_copy_TOF_GUMS_{code}.json"
        if not src.exists():
            print(f"[{code}] no translation file")
            continue
        data = json.loads(src.read_text())
        lang_folder = FOLDER / folder_name
        if not lang_folder.exists():
            print(f"[{code}] no image folder: {lang_folder}")
            continue
        slug = name.replace(" ", "_").replace("(", "").replace(")", "")
        out = FOLDER / f"GUMS_TOF_{code}_{slug}.docx"
        ads, skipped = build_doc(code, name, lang_folder, markets,
                                  f"inaessentials.{folder_name.lower()}",
                                  data, THEME_MAP, out)
        size_kb = out.stat().st_size // 1024
        print(f"[{code}] wrote {out.name} ({size_kb} KB, {ads} ads, {skipped} missing images)")


if __name__ == "__main__":
    main()
