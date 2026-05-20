"""
Build Bulgarian ad copy .docx files for the two AI ADS MAY/ subfolders.
No API calls. All copy is authored inline below.

Brand: Ina Essentials, Хидролина Скумпия (Smoke Tree Hydrolina oral spray).
Voice: Bulgarian Cyrillic, informal "ти", NO em-dashes / en-dashes.
Format per creative: 3 headlines + medium primary text (3-4 sentences) + CTA.
"""
import io
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

ROOT = Path("/Users/magi/Downloads/AI ADS MAY")
IMG_MAX_DIM_PX = 1200
IMG_MAX_WIDTH_INCHES = 5.0


# ════════════════════════════════════════════════════════════════════════════
#   GENERAL FOLDER (top of funnel, broader hooks, 12 creatives)
# ════════════════════════════════════════════════════════════════════════════
GENERAL = [
    {
        "filename": "01_breaking_news_gum_headline.png",
        "angle": "Heritage / forgotten Bulgarian remedy. Faux magazine cover anchors authority and nostalgia.",
        "headlines": [
            "Българската билка, която баба ползваше",
            "Преди пастата имаше Скумпия",
            "Старият навик за здрави венци",
        ],
        "primary": (
            "Преди да съществува пастата за зъби, в българските села е имало друг начин за здрави венци "
            "и свеж дъх. Билка със смрадлив корен, Скумпия, която баячките са ползвали за венците. "
            "Днес тя се връща в малък флакон спрей. Хидролина Скумпия събира този забравен ритуал в едно движение."
        ),
        "cta": "Научи повече",
    },
    {
        "filename": "03_mystery_hook_article.png",
        "angle": "Authority / mystery. Doctor's notepad style implies clinical insight without claiming it.",
        "headlines": [
            "Една причина, три симптома",
            "Защо венците ти кървят",
            "Източникът е в устата ти",
        ],
        "primary": (
            "Кървене, възпаление, лош дъх: изглеждат различни проблеми, но имат един и същи корен. "
            "Микробиомът в устата е изваден от равновесие. Българската скумпия от векове помага за "
            "връщането на този баланс. Хидролина Скумпия пренася тази билка в спрей за всеки ден."
        ),
        "cta": "Открий продукта",
    },
    {
        "filename": "04_whiteboard_why_bad_breath_returns.png",
        "angle": "Education. Why does it keep returning? Reframes the problem so the product is the obvious answer.",
        "headlines": [
            "Защо лошият дъх се връща",
            "Не е виновна пастата",
            "Истинската причина за дъха",
        ],
        "primary": (
            "Миеш зъби, дъвчеш дъвка, и след час дъхът е същият. Защото пастата не стига до езика "
            "и до бактериите близо до венците. Скумпията е българска билка, която работи там, "
            "където четката не достига. Хидролина Скумпия е спрей за устната грижа, която ти липсваше."
        ),
        "cta": "Научи повече",
    },
    {
        "filename": "05_anatomy_diagram_tongue_bacteria.png",
        "angle": "Educational anatomy. Names the real culprit (the tongue) and positions the spray as the only thing that reaches it.",
        "headlines": [
            "80% от дъха започва от езика",
            "Не зъбите. Езикът.",
            "Бактериите живеят тук",
        ],
        "primary": (
            "Само 20% от лошия дъх идва от зъбите. Останалите 80% са от бактерии по задната част на езика. "
            "Скумпията, българската билка, е традиционно средство точно за тази зона. "
            "Хидролина Скумпия е спрей, който влиза там, където четката не стига."
        ),
        "cta": "Виж повече",
    },
    {
        "filename": "07_stat_surround_one_in_three.png",
        "angle": "Stat hook. Wide statistic makes the problem feel common and undiagnosed.",
        "headlines": [
            "1 от 3 има проблем с венците",
            "Тиха статистика, силен сигнал",
            "Ти може да си третият",
        ],
        "primary": (
            "Един от трима възрастни има възпалени или кървящи венци, без да го забелязва. "
            "Симптомите идват тихо: розово в мивката, лош дъх, чувствителност. Скумпията е "
            "българска билка, използвана от векове за устната грижа. Хидролина Скумпия я връща в спрей за всеки ден."
        ),
        "cta": "Научи повече",
    },
    {
        "filename": "08_ugc_mirror_teeth_check.png",
        "angle": "Self-diagnosis curiosity. Invites the viewer to do the test now, in the mirror.",
        "headlines": [
            "Виж какво показва огледалото",
            "Венците ти казват истината",
            "Огледален тест за венци",
        ],
        "primary": (
            "Дръпни леко устната пред огледалото. Венците розови или тъмночервени? Меки или подути? "
            "Кървят ли при четкане? Това са най-ранните сигнали, които повечето хора пропускат. "
            "Хидролина Скумпия е спрей с българската билка, която традиционно се ползва точно за тези признаци."
        ),
        "cta": "Виж повече",
    },
    {
        "filename": "09_ugc_hand_covering_mouth.png",
        "angle": "Social shame. Reframes embarrassment into agency: stop hiding, start solving.",
        "headlines": [
            "Не покривай. Решавай.",
            "Свободно дишане, без срам",
            "Дъхът не се крие",
        ],
        "primary": (
            "Покриваш устата, когато се смееш. Дъвчеш дъвка преди срещата. Маскираш и пак същото. "
            "Защото проблемът не е по емайла, а в езика, венците и слюнката. Хидролина Скумпия е "
            "български спрей със скумпия, билката, която работи там, където дъвките не стигат."
        ),
        "cta": "Купи сега",
    },
    {
        "filename": "10_ugc_toothbrush_moment.png",
        "angle": "Relatable problem moment. Names the pink in the sink without sounding alarming.",
        "headlines": [
            "Розово в мивката? Не пренебрегвай.",
            "Не е „просто така“",
            "Първият сигнал не е финалът",
        ],
        "primary": (
            "Виждаш малко розово в мивката, докато миеш зъби. Не е от четката. Кървящите венци са "
            "най-ранният сигнал, че микробиомът в устата ти е извън баланс. Българската скумпия от "
            "векове е традиционно средство за тази зона. Хидролина Скумпия я връща в спрей формат."
        ),
        "cta": "Научи повече",
    },
    {
        "filename": "11_pink_smoke_tree_with_headline.png",
        "angle": "Beauty plus heritage. The pink bloom is the hero; copy explains why this herb matters.",
        "headlines": [
            "Розова билка от Балкана",
            "Българска грижа за устата",
            "Скумпията не цъфти случайно",
        ],
        "primary": (
            "Това е Cotinus coggygria. Българите я наричат Скумпия. Цъфти розово по балканските "
            "склонове и от векове е част от народната медицина за венците и дъха. Хидролина Скумпия "
            "пренася тази билка в малък флакон спрей."
        ),
        "cta": "Открий продукта",
    },
    {
        "filename": "13_leaf_macro_with_stat.png",
        "angle": "Educational beauty. Names what's inside the leaf so the visual carries scientific weight.",
        "headlines": [
            "Бордо лист, силна билка",
            "Малък лист, голям ефект",
            "Скумпията под микроскоп",
        ],
        "primary": (
            "Бордовите овални листа на скумпията крият активни вещества, които българите познават "
            "от векове. Танини, флавоноиди, полифеноли. Затова народът е посягал към тази билка "
            "за венците и дъха. Хидролина Скумпия събира тази сила в едно пръскане."
        ),
        "cta": "Виж повече",
    },
    {
        "filename": "14_color_block_bold_statement.png",
        "angle": "Provocative typographic statement. Disqualifies the wrong solution (toothpaste alone) and offers the right one.",
        "headlines": [
            "Лош дъх не се измива",
            "Пастата не е решението",
            "Истината за устната грижа",
        ],
        "primary": (
            "Можеш да миеш зъбите три пъти на ден и пак да имаш лош дъх. Защото проблемът не е по "
            "емайла, а по езика, във венците и в слюнката. Българската скумпия е традиционно решение "
            "за точно тези зони. Хидролина Скумпия е спрей, който мие там, където четката пропуска."
        ),
        "cta": "Научи повече",
    },
    {
        "filename": "A05_tongue_coating.png",
        "angle": "Tongue coating diagnostic. Image already says 80% of bad breath starts here; copy explains why.",
        "headlines": [
            "Налепът тежи по-силно от чесъна",
            "Чист език, свеж дъх",
            "Истинският източник на дъха",
        ],
        "primary": (
            "Бялото покритие на задната част на езика не е безобидно. Това са бактерии, които "
            "произвеждат серни съединения, основната причина за лошия дъх. Скумпията е българска "
            "билка, която традиционно се ползва за устната грижа. Хидролина Скумпия е спрей с тази "
            "билка, който достига там, където четката не отива."
        ),
        "cta": "Открий продукта",
    },
]


# ════════════════════════════════════════════════════════════════════════════
#   GUM RECESSION FOLDER (symptom-aware diagnostic, 12 creatives)
# ════════════════════════════════════════════════════════════════════════════
GUM_RECESSION = [
    {
        "filename": "A01_gum_recession_diagnostic.png",
        "angle": "Recession diagnostic. Image already labels recession; copy explains the irreversibility and offers the herbal route.",
        "headlines": [
            "Венците не растат обратно",
            "Българската билка за венците",
            "Спри причината, не симптома",
        ],
        "primary": (
            "Зъбите ти изглеждат по-дълги от преди. Венците са се прибрали и са оголили основата на зъба. "
            "Това е рецесия и не се връща сама. Българската скумпия е традиционна билка за грижа за венците. "
            "Хидролина Скумпия я носи в спрей, който достига до самата граница на венеца."
        ),
        "cta": "Научи повече",
    },
    {
        "filename": "A02_bleeding_gums_diagnostic.png",
        "angle": "Bleeding gums diagnostic. The pink-in-sink moment, reframed as a first signal worth listening to.",
        "headlines": [
            "Розово в мивката, ранен сигнал",
            "Кървенето не е нормално",
            "Слушай венците си",
        ],
        "primary": (
            "Малко розово в мивката, докато миеш зъби, не е „просто така“. Кървящите венци са "
            "първият тих сигнал, че микробиомът в устата ти е извън баланс. Българите от векове "
            "посягат към една билка за това: скумпията. Хидролина Скумпия е спрей за устна грижа, "
            "който връща този стар навик в днешен формат."
        ),
        "cta": "Открий продукта",
    },
    {
        "filename": "A03_gingivitis_inflammation.png",
        "angle": "Gingivitis diagnostic. Calls out swollen red gums as stage one, not stage zero.",
        "headlines": [
            "Червеното е сигнал, не цвят",
            "Подути венци не са нормални",
            "Първият стадий е тих",
        ],
        "primary": (
            "Здравите венци са нежно розови и плътни. Подути, тъмночервени, болезнени при докосване? "
            "Това е възпаление, първият стадий преди по-сериозните проблеми. Скумпията е българска "
            "билка с традиционна употреба за венците. Хидролина Скумпия я носи в спрей за всекидневна грижа."
        ),
        "cta": "Научи повече",
    },
    {
        "filename": "A04_tartar_at_gumline.png",
        "angle": "Tartar diagnostic. Re-attributes the yellow stripe (not coffee, mineralized plaque).",
        "headlines": [
            "Жълтото не е от кафето",
            "Зъбният камък е минерал",
            "Под видимото има друго",
        ],
        "primary": (
            "Жълтите ивици в основата на зъбите изглеждат като петна от кафе или вино. Но не са. "
            "Това е минерализирана плака, зъбен камък, в който живеят бактерии. Скумпията е "
            "българска билка, която традиционно се ползва за устната грижа. Хидролина Скумпия е "
            "спрей, който достига точно тази зона."
        ),
        "cta": "Виж повече",
    },
    {
        "filename": "A06_exposed_tooth_root.png",
        "angle": "Exposed root diagnostic. Visual cue (longer teeth) becomes the hook.",
        "headlines": [
            "Зъбите не порастват сами",
            "Венецът се прибира тихо",
            "Дължина, която не искаш",
        ],
        "primary": (
            "Зъбите ти изглеждат по-дълги? Не са, венецът се е прибрал и е оголил корена. Този процес "
            "е тих и постепенен, и не се обръща сам. Скумпията е традиционна българска билка за "
            "венците. Хидролина Скумпия е спрей с тази билка, който се ползва всекидневно."
        ),
        "cta": "Научи повече",
    },
    {
        "filename": "A07_periodontal_pocket.png",
        "angle": "Periodontal pocket diagnostic. Names the hidden space where the brush never reaches.",
        "headlines": [
            "Тих джоб, силни последици",
            "Между зъб и венец",
            "Скритата зона на устата",
        ],
        "primary": (
            "Дълбокият джоб между зъба и венеца е място, до което четката не стига. Там бактериите "
            "живеят необезпокоявани, хранят се с плака и възпаляват тъканта. Българската скумпия е "
            "традиционна билка за устната грижа. Хидролина Скумпия е спрей, който влиза точно в тази зона."
        ),
        "cta": "Открий продукта",
    },
    {
        "filename": "A08_plaque_buildup.png",
        "angle": "Plaque diagnostic. Names what the mirror can't show you.",
        "headlines": [
            "Невидимата плака",
            "Това, което четката пропуска",
            "Тънкият филм по емайла",
        ],
        "primary": (
            "Тънкият жълтеникав филм по емайла е невидим в огледалото. Но точно той е началото на "
            "зъбния камък, кървенето и възпалението. Скумпията е българска билка, която традиционно "
            "се ползва за устната грижа. Хидролина Скумпия е спрей, който мие там, където четката пропуска."
        ),
        "cta": "Виж повече",
    },
    {
        "filename": "hf_20260514_142022_79e54613-3c54-4a3e-b21a-7362b8ec9547.png",
        "angle": "Herb hero shot. Lets the visual carry the brand and the copy do the introduction.",
        "headlines": [
            "Това е Скумпия",
            "Билката, която българите познават",
            "Розов цвят, силна история",
        ],
        "primary": (
            "Розовите облаци на скумпията не са случайни. Балканите дават на тази билка особена сила, "
            "която българите познават от векове. Народът я ползва за венците, за дъха, за устната грижа. "
            "Хидролина Скумпия я носи в малък флакон спрей."
        ),
        "cta": "Открий продукта",
    },
    {
        "filename": "hf_20260514_142124_10609a2d-6875-4003-bc2b-749cd629f0d4.png",
        "angle": "Tradition meets modern format. Old recipe, new vehicle.",
        "headlines": [
            "Балканска билка за устата",
            "Скумпия в спрей формат",
            "Стара рецепта, нов вид",
        ],
        "primary": (
            "Векове наред българите ползват скумпията за здрави венци и свеж дъх. Днес тази билка се "
            "връща в модерен формат: спрей за устна грижа. Без сложен ритуал, без посещения. Едно "
            "пръскане, ден или вечер, и забравената билка е отново в устата ти."
        ),
        "cta": "Купи сега",
    },
    {
        "filename": "hf_20260515_060236_d5319956-8e84-4edf-95c2-91c8162495a9.png",
        "angle": "Morning hook. The driest hour of the day is where the spray earns its place.",
        "headlines": [
            "Свеж дъх от Балкана",
            "Сутрин с билката",
            "Утринен ритуал, древна билка",
        ],
        "primary": (
            "Сутрин устата е суха, бактериите са празнували цяла нощ, дъхът е тежък. Българите от векове "
            "посягат към една билка за това: скумпията. Хидролина Скумпия е спрей, който поема тази "
            "билка и я носи навсякъде. Едно пръскане сутрин, и устата се събужда."
        ),
        "cta": "Виж повече",
    },
    {
        "filename": "hf_20260515_060324_c4173df7-fec4-45b6-a768-7192c042feb1.png",
        "angle": "Heritage credibility. The herb is not a trend, it's a tradition.",
        "headlines": [
            "От Балкана за венците",
            "Билка с тиха сила",
            "Корени в българската земя",
        ],
        "primary": (
            "Скумпията расте по Балканите от хиляди години. Българските баячки са я ползвали за венците, "
            "за дъха, за раничките в устата. Не защото е модерна, а защото работи. Хидролина Скумпия "
            "връща тази билка в спрей за днешния ден."
        ),
        "cta": "Научи повече",
    },
    {
        "filename": "hf_20260515_060350_10be5a72-7baa-46d4-b014-c82b652cb80d.png",
        "angle": "Closing pitch. Simple, calm, evergreen.",
        "headlines": [
            "Старата българска грижа",
            "Скумпия в малък флакон",
            "Билка, която работи",
        ],
        "primary": (
            "Преди да съществуват сините марки и белите тубички, в българските села е имало друго: "
            "скумпията. Билка със смрадлив корен и силно действие за венците. Днес тя се връща в спрей "
            "формат, удобен за чанта или нощно шкафче. Хидролина Скумпия е този ритуал, опростен."
        ),
        "cta": "Открий продукта",
    },
]


# ════════════════════════════════════════════════════════════════════════════
#   Image embedding helper
# ════════════════════════════════════════════════════════════════════════════
def prepare_image_bytes(src_path: Path) -> io.BytesIO:
    with Image.open(src_path) as im:
        im.load()
        if im.mode in ("RGBA", "P"):
            bg = Image.new("RGB", im.size, (255, 255, 255))
            if im.mode == "P":
                im = im.convert("RGBA")
            bg.paste(im, mask=im.split()[-1] if im.mode == "RGBA" else None)
            im = bg
        elif im.mode != "RGB":
            im = im.convert("RGB")
        w, h = im.size
        scale = min(1.0, IMG_MAX_DIM_PX / max(w, h))
        if scale < 1.0:
            im = im.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
        buf = io.BytesIO()
        im.save(buf, format="JPEG", quality=82, optimize=True)
        buf.seek(0)
        return buf


def strip_dashes(text: str) -> str:
    if not text:
        return text
    return text.replace("—", ", ").replace("–", ", ")


# ════════════════════════════════════════════════════════════════════════════
#   Doc builder
# ════════════════════════════════════════════════════════════════════════════
def build_docx(folder: Path, docx_name: str, folder_desc: str, items: list[dict]):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Ina Essentials, Хидролина Скумпия", level=0)
    doc.add_heading(f"Ad copy, Bulgarian, {folder.name}", level=1)

    p = doc.add_paragraph()
    p.add_run("Папка: ").bold = True
    p.add_run(folder_desc)

    p = doc.add_paragraph()
    p.add_run("Общо реклами: ").bold = True
    p.add_run(str(len(items)))

    p = doc.add_paragraph()
    p.add_run("Формат: ").bold = True
    p.add_run("3 заглавия (опции), основен текст, CTA. Без тирета. Език: български.")

    doc.add_paragraph()

    for r in items:
        doc.add_page_break()
        doc.add_heading(r["filename"], level=2)

        img_path = folder / r["filename"]
        if img_path.exists():
            try:
                buf = prepare_image_bytes(img_path)
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(buf, width=Inches(IMG_MAX_WIDTH_INCHES))
            except Exception as e:
                doc.add_paragraph(f"[image error: {e}]")
        else:
            doc.add_paragraph(f"[image not found: {r['filename']}]")

        if r.get("angle"):
            p = doc.add_paragraph()
            p.add_run("Ad Angle: ").bold = True
            p.add_run(r["angle"])

        p = doc.add_paragraph()
        p.add_run("Заглавия (3 опции):").bold = True
        for i, h in enumerate(r["headlines"], 1):
            doc.add_paragraph(f"{i}. {strip_dashes(h)}")

        p = doc.add_paragraph()
        p.add_run("Основен текст:").bold = True
        doc.add_paragraph(strip_dashes(r["primary"]))

        p = doc.add_paragraph()
        p.add_run("CTA: ").bold = True
        p.add_run(strip_dashes(r["cta"]))

        sep = doc.add_paragraph("_" * 50)
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

    out = folder / docx_name
    doc.save(str(out))
    return out


def main():
    pairs = [
        (ROOT / "General", "General_AdCopy_BG.docx",
         "General creatives, broader top-of-funnel hooks", GENERAL),
        (ROOT / "New Concept (GUM RECCESION)", "GumRecession_AdCopy_BG.docx",
         "Diagnostic close-ups: gum recession, bleeding gums, gingivitis, tartar, plaque, periodontal pocket, smoke tree heritage shots",
         GUM_RECESSION),
    ]
    for folder, name, desc, items in pairs:
        if not folder.is_dir():
            print(f"[SKIP] not a directory: {folder}")
            continue
        out = build_docx(folder, name, desc, items)
        kb = out.stat().st_size // 1024
        print(f"wrote {out} ({kb} KB, {len(items)} ads)")


if __name__ == "__main__":
    main()
