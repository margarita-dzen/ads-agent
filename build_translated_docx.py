"""
Build CZ / SK / GR .docx and JSON copies of the Smoke Tree Hydrolina ads,
mirroring the BG build. .docx for human review, JSON for the ads-uploader CLI.
"""
import io
import json
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

from translations_skumpia import LANGUAGES  # CZ, SK, GR

ROOT = Path("/Users/magi/Downloads/AI ADS MAY")
IMG_MAX_DIM_PX = 1200
IMG_MAX_WIDTH_INCHES = 5.0

FOLDERS = [
    ("General", "General"),                                   # subdir, friendly label
    ("New Concept (GUM RECCESION)", "GumRecession"),
]


def prepare_image_bytes(src_path: Path) -> io.BytesIO:
    with Image.open(src_path) as im:
        im.load()
        if im.mode in ("RGBA", "P"):
            bg = Image.new("RGB", im.size, (255, 255, 255))
            if im.mode == "P":
                im = im.convert("RGBA")
            bg.paste(im, mask=im.split()[-1] if im.mode == "RGBA" else None)
            im = bg
        elif im.mode != "RGB":
            im = im.convert("RGB")
        w, h = im.size
        scale = min(1.0, IMG_MAX_DIM_PX / max(w, h))
        if scale < 1.0:
            im = im.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
        buf = io.BytesIO()
        im.save(buf, format="JPEG", quality=82, optimize=True)
        buf.seek(0)
        return buf


def strip_dashes(text: str) -> str:
    return text.replace("—", ", ").replace("–", ", ")


LABELS = {
    "CZ": dict(folder="Složka", total="Celkem reklam", fmt="Formát",
               fmt_value="3 nadpisy (varianty), hlavní text, CTA. Bez pomlček. Jazyk: čeština.",
               doc_title="Ad copy, čeština", angle="Ad Angle",
               heads="Nadpisy (3 varianty):", primary="Hlavní text:", cta="CTA"),
    "SK": dict(folder="Priečinok", total="Celkom reklám", fmt="Formát",
               fmt_value="3 nadpisy (varianty), hlavný text, CTA. Bez pomlčiek. Jazyk: slovenčina.",
               doc_title="Ad copy, slovenčina", angle="Ad Angle",
               heads="Nadpisy (3 varianty):", primary="Hlavný text:", cta="CTA"),
    "GR": dict(folder="Φάκελος", total="Συνολικές διαφημίσεις", fmt="Φόρμα",
               fmt_value="3 τίτλοι (εναλλακτικές), κύριο κείμενο, CTA. Χωρίς παύλες. Γλώσσα: ελληνικά.",
               doc_title="Ad copy, ελληνικά", angle="Ad Angle",
               heads="Τίτλοι (3 εναλλακτικές):", primary="Κύριο κείμενο:", cta="CTA"),
}


def build_docx(folder: Path, out_path: Path, items: list[dict], lang: str, folder_desc: str):
    L = LABELS[lang]
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Ina Essentials, Hydrolina Skumpia", level=0)
    doc.add_heading(f"{L['doc_title']}, {folder.name}", level=1)

    p = doc.add_paragraph()
    p.add_run(f"{L['folder']}: ").bold = True
    p.add_run(folder_desc)

    p = doc.add_paragraph()
    p.add_run(f"{L['total']}: ").bold = True
    p.add_run(str(len(items)))

    p = doc.add_paragraph()
    p.add_run(f"{L['fmt']}: ").bold = True
    p.add_run(L["fmt_value"])

    doc.add_paragraph()

    for r in items:
        doc.add_page_break()
        doc.add_heading(r["filename"], level=2)

        img_path = folder / r["filename"]
        if img_path.exists():
            try:
                buf = prepare_image_bytes(img_path)
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(buf, width=Inches(IMG_MAX_WIDTH_INCHES))
            except Exception as e:
                doc.add_paragraph(f"[image error: {e}]")
        else:
            doc.add_paragraph(f"[image not found: {r['filename']}]")

        if r.get("angle"):
            p = doc.add_paragraph()
            p.add_run(f"{L['angle']}: ").bold = True
            p.add_run(r["angle"])

        p = doc.add_paragraph()
        p.add_run(L["heads"]).bold = True
        for i, h in enumerate(r["headlines"], 1):
            doc.add_paragraph(f"{i}. {strip_dashes(h)}")

        p = doc.add_paragraph()
        p.add_run(L["primary"]).bold = True
        doc.add_paragraph(strip_dashes(r["primary"]))

        p = doc.add_paragraph()
        p.add_run(f"{L['cta']}: ").bold = True
        p.add_run(strip_dashes(r["cta"]))

        sep = doc.add_paragraph("_" * 50)
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(out_path))


def main():
    for lang_code, (lang_name, gen_items, gum_items) in LANGUAGES.items():
        for (subdir, label), items in zip(FOLDERS, [gen_items, gum_items]):
            folder = ROOT / subdir
            if not folder.is_dir():
                print(f"[SKIP] missing folder {folder}")
                continue

            docx_path = folder / f"{label}_AdCopy_{lang_code}.docx"
            build_docx(folder, docx_path, items, lang_code,
                       folder_desc=f"{subdir} ({lang_name})")
            print(f"wrote {docx_path}  ({docx_path.stat().st_size // 1024} KB)")

            # JSON twin for the ads-uploader CLI (no images embedded, just text)
            json_path = folder / f"{label}_AdCopy_{lang_code}.json"
            json_path.write_text(
                json.dumps(items, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            print(f"wrote {json_path}  ({json_path.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
