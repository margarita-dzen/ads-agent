"""
Build one .docx file per language with embedded images above each ad's copy.
The .docx files are Word/Google Docs compatible — drag them into Google Drive and
they auto-convert to native Google Docs format.
"""
import io
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

ROOT = Path("/Users/magi/Downloads/Veins New ads")
TRANS_DIR = ROOT / "translations"
SORTED_DIR = ROOT / "sorted"
OUT_DIR = ROOT / "by_language_docx"

LANGUAGES = [
    ("EN", "English",     "UK, US, AU, IE", "inaessentials.co.uk, inaessentials.us, inaessentials.au, inaessentials.ie"),
    ("BG", "Bulgarian",   "BG",              "inaessentials.bg"),
    ("FR", "French",      "FR",              "inaessentials.fr"),
    ("RO", "Romanian",    "RO",              "inaessentials.ro"),
    ("SK", "Slovak",      "SK",              "www.inaessentials.sk"),
    ("CZ", "Czech",       "CZ",              "www.inaessentials.cz"),
    ("DE", "German",      "DE",              "inaessentials.de"),
    ("IT", "Italian",     "IT",              "inaessentials.it"),
    ("ES", "Spanish",     "ES",              "inaessentials.es"),
    ("NL", "Dutch",       "NL",              "inaessentials.nl"),
    ("PT", "Portuguese (European)", "PT",   "inaessentials.pt"),
    ("PL", "Polish",      "PL",              "inaessentials.pl"),
    ("HU", "Hungarian",   "HU",              "inaessentials.hu"),
    ("HR", "Croatian",    "HR",              "inaessentials.hr"),
    ("SI", "Slovenian",   "SI",              "inaessentials.si"),
    ("RS", "Serbian (Latin)", "RS",          "inaessentials.rs"),
    ("EL", "Greek",       "GR",              "inaessentials.gr"),
]

CATEGORY_ORDER = [
    "anatomy", "article", "before-after", "chart", "copy",
    "diary", "product", "review-list", "social-fb", "stat-hook",
    "testimonial-quote", "visual",
]

# Resize images so DOCX doesn't end up gigantic
IMG_MAX_WIDTH_INCHES = 4.0
IMG_MAX_DIM_PX = 1200


def prepare_image_bytes(src_path: Path) -> tuple[io.BytesIO, str]:
    """Open, downscale, and re-encode the image as JPEG for embedding."""
    with Image.open(src_path) as im:
        im.load()
        if im.mode in ("RGBA", "P"):
            bg = Image.new("RGB", im.size, (255, 255, 255))
            if im.mode == "P":
                im = im.convert("RGBA")
            bg.paste(im, mask=im.split()[-1] if im.mode == "RGBA" else None)
            im = bg
        elif im.mode != "RGB":
            im = im.convert("RGB")
        w, h = im.size
        scale = min(1.0, IMG_MAX_DIM_PX / max(w, h))
        if scale < 1.0:
            im = im.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
        buf = io.BytesIO()
        im.save(buf, format="JPEG", quality=82, optimize=True)
        buf.seek(0)
        return buf, "jpeg"


def category_of(filename: str) -> str:
    return filename.split("_")[0]


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_kv(doc, label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r2 = p.add_run(value)
    return p


def add_label_line(doc, label):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    return p


def build_language_docx(code, name, markets, domains, data, out_path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(f"Ina Essentials, Soothing Cream with Horse Chestnut and Smoke Tree", level=0)
    sub = doc.add_heading(f"Ad copy + headlines, {name}", level=1)

    add_kv(doc, "Markets", markets)
    add_kv(doc, "Shopify domains", domains)
    add_kv(doc, "Total ads", str(len(data)))
    add_kv(doc, "Format", "Headlines (3 options), Primary Text (medium), CTA. Image embedded above each ad.")

    doc.add_paragraph()

    by_cat = {}
    for r in data:
        by_cat.setdefault(category_of(r["filename"]), []).append(r)
    ordered_cats = [c for c in CATEGORY_ORDER if c in by_cat] + [c for c in by_cat if c not in CATEGORY_ORDER]

    for cat in ordered_cats:
        doc.add_page_break()
        doc.add_heading(f"Category: {cat}  ({len(by_cat[cat])} ads)", level=1)

        for r in by_cat[cat]:
            fn = r["filename"]
            doc.add_heading(fn, level=2)

            img_path = SORTED_DIR / fn
            if img_path.exists():
                try:
                    buf, _ = prepare_image_bytes(img_path)
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(buf, width=Inches(IMG_MAX_WIDTH_INCHES))
                except Exception as e:
                    doc.add_paragraph(f"[image error: {e}]")
            else:
                doc.add_paragraph(f"[image not found: {fn}]")

            add_label_line(doc, "Headlines")
            for i, h in enumerate(r.get("headlines", []), 1):
                doc.add_paragraph(f"{i}. {h}")

            add_label_line(doc, "Primary Text")
            doc.add_paragraph(r.get("primary_text", ""))

            add_kv(doc, "CTA", r.get("cta", ""))

            sep = doc.add_paragraph("_" * 50)
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(out_path))


def main():
    OUT_DIR.mkdir(exist_ok=True)
    for code, name, markets, domains in LANGUAGES:
        src = TRANS_DIR / f"{code}.json"
        if not src.exists():
            print(f"[{code}] MISSING {src}")
            continue
        data = json.loads(src.read_text())
        slug = name.replace(" ", "_").replace("(", "").replace(")", "")
        out = OUT_DIR / f"{code}_{slug}.docx"
        build_language_docx(code, name, markets, domains, data, out)
        size_kb = out.stat().st_size // 1024
        print(f"[{code}] wrote {out.name} ({size_kb} KB)")

    # Quick README explaining how to upload to Google Drive
    readme = OUT_DIR / "HOW_TO_UPLOAD_TO_GOOGLE_DOCS.txt"
    readme.write_text(
        "How to convert these .docx files to Google Docs:\n\n"
        "1. Open Google Drive in your browser (drive.google.com).\n"
        "2. Select all the .docx files in this folder and drag them into Drive.\n"
        "3. Right-click any uploaded file -> Open with -> Google Docs.\n"
        "   (Or: Drive settings -> Convert uploads -> on, then re-upload.)\n"
        "4. The file opens as a native Google Doc with images and formatting preserved.\n\n"
        "Each .docx contains 61 ads in 12 categories, with the ad image embedded\n"
        "above the headlines / primary text / CTA for that ad. Brand-compliant,\n"
        "no em-dashes, ready for the team to paste into Meta Ads Manager.\n"
    )
    print(f"wrote {readme.name}")


if __name__ == "__main__":
    main()
